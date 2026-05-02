from fastapi import FastAPI, HTTPException
import hashlib
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List
import anthropic
import os
import io

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["POST", "GET", "DELETE", "PATCH"],
    allow_headers=["*"],
)

client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

# ── PostgreSQL ─────────────────────────────────────────────────────────────────
DATABASE_URL = os.environ.get("DATABASE_URL", "")
db_pool = None

@app.on_event("startup")
async def startup():
    global db_pool
    if DATABASE_URL:
        try:
            from psycopg_pool import AsyncConnectionPool
            url = DATABASE_URL.replace("postgres://", "postgresql://")
            db_pool = AsyncConnectionPool(url, min_size=1, max_size=5, open=False)
            await db_pool.open()
            await init_db()
            print("PostgreSQL verbunden")
        except Exception as e:
            print(f"DB Verbindung fehlgeschlagen: {e} — nutze In-Memory")
            db_pool = None
    else:
        print("Kein DATABASE_URL — nutze In-Memory")

@app.on_event("shutdown")
async def shutdown():
    if db_pool:
        await db_pool.close()

async def init_db():
    async with db_pool.connection() as conn:
        await conn.execute("""
            CREATE TABLE IF NOT EXISTS otp_sessions (
                session_id    TEXT PRIMARY KEY,
                phone         TEXT NOT NULL,
                name          TEXT NOT NULL,
                dob           TEXT NOT NULL,
                code          TEXT NOT NULL,
                auth_token    TEXT,
                verified      BOOLEAN DEFAULT FALSE,
                attempts      INTEGER DEFAULT 0,
                expires_at    TIMESTAMPTZ NOT NULL,
                verified_at   TIMESTAMPTZ,
                created_at    TIMESTAMPTZ DEFAULT NOW()
            )
        """)
        await conn.execute("""
            CREATE TABLE IF NOT EXISTS appointments (
                id             SERIAL PRIMARY KEY,
                session_id     TEXT NOT NULL,
                patient_name   TEXT NOT NULL,
                patient_phone  TEXT NOT NULL,
                patient_dob    TEXT,
                appointment_id TEXT NOT NULL,
                therapeut      TEXT,
                standort       TEXT,
                datum          TEXT,
                zeit           TEXT,
                datetime_iso   TEXT,
                entity_id      TEXT,
                prof_id        TEXT,
                calendar_id    TEXT,
                status         TEXT DEFAULT 'active',
                gebucht_am     TIMESTAMPTZ DEFAULT NOW(),
                abgesagt_am    TIMESTAMPTZ
            )
        """)
        await conn.execute(
            "CREATE INDEX IF NOT EXISTS idx_app_session ON appointments(session_id)"
        )
        await conn.execute(
            "CREATE INDEX IF NOT EXISTS idx_app_phone ON appointments(patient_phone)"
        )

async def db_save_otp(session_id, phone, name, dob, code, expires):
    if not db_pool: return
    import datetime
    exp = datetime.datetime.fromtimestamp(expires, tz=datetime.timezone.utc)
    async with db_pool.connection() as conn:
        await conn.execute("""
            INSERT INTO otp_sessions (session_id, phone, name, dob, code, expires_at)
            VALUES (%s,%s,%s,%s,%s,%s)
            ON CONFLICT (session_id) DO UPDATE
            SET code=EXCLUDED.code, expires_at=EXCLUDED.expires_at,
                verified=FALSE, attempts=0
        """, (session_id, phone, name, dob, code, exp))

async def db_verify_otp(session_id, code):
    if not db_pool: return None
    import datetime, secrets as _s
    async with db_pool.connection() as conn:
        async with conn.cursor(row_factory=__import__('psycopg').rows.dict_row) as cur:
            await cur.execute(
                "SELECT * FROM otp_sessions WHERE session_id=%s", (session_id,)
            )
            row = await cur.fetchone()
        if not row: return None
        if datetime.datetime.now(tz=datetime.timezone.utc) > row["expires_at"]:
            return {"error": "abgelaufen"}
        if row["attempts"] >= 3:
            return {"error": "zu_viele_versuche"}
        await conn.execute(
            "UPDATE otp_sessions SET attempts=attempts+1 WHERE session_id=%s", (session_id,)
        )
        if row["code"] != code:
            return {"error": "falsch"}
        token = _s.token_urlsafe(24)
        now = datetime.datetime.now(tz=datetime.timezone.utc)
        await conn.execute("""
            UPDATE otp_sessions SET verified=TRUE, auth_token=%s, verified_at=%s
            WHERE session_id=%s
        """, (token, now, session_id))
        return {"auth_token": token, "name": row["name"], "dob": row["dob"], "phone": row["phone"]}

async def db_save_appointment(session_id, data):
    if not db_pool: return
    async with db_pool.connection() as conn:
        await conn.execute("""
            INSERT INTO appointments
            (session_id, patient_name, patient_phone, patient_dob,
             appointment_id, therapeut, standort, datum, zeit, datetime_iso,
             entity_id, prof_id, calendar_id)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
        """, (
            session_id,
            data.get("patient_name",""), data.get("patient_phone",""), data.get("patient_dob",""),
            str(data.get("appointment_id","")),
            data.get("therapeut",""), data.get("standort",""),
            data.get("datum",""), data.get("zeit",""), data.get("datetime_iso",""),
            data.get("entity_id",""), data.get("prof_id",""), data.get("calendar_id",""),
        ))

async def db_get_appointments(session_id, phone):
    if not db_pool: return []
    async with db_pool.connection() as conn:
        async with conn.cursor(row_factory=__import__('psycopg').rows.dict_row) as cur:
            await cur.execute("""
                SELECT * FROM appointments
                WHERE (session_id=%s OR patient_phone=%s) AND status='active'
                ORDER BY datum ASC, zeit ASC
            """, (session_id, phone))
            rows = await cur.fetchall()
            return [dict(r) for r in rows]

async def db_cancel_appointment(appointment_id, session_id, phone):
    if not db_pool: return False
    async with db_pool.connection() as conn:
        cur = await conn.execute("""
            UPDATE appointments SET status='cancelled', abgesagt_am=NOW()
            WHERE appointment_id=%s
            AND (session_id=%s OR patient_phone=%s)
            AND status='active'
        """, (appointment_id, session_id, phone))
        return cur.rowcount == 1

SYSTEM_PROMPT = """
Du bist der Kineo Assistent — digitaler Begleiter von Kineo AG, Funktionelle Physiotherapie, 6 Standorte Zürich.

=== NUTZER-TYPEN ===
Standard: Patient. Begleite ihn: Symptome → Therapeut-Empfehlung → Standort → Termine laden → Buchen.
Team: Gibt bekannten Namen ein → frage nach Team-PIN (${TEAM_PIN}) → bei richtigem PIN: Stellenplan, Zimmer, HR-Infos zugänglich.
Manager (intern, NIE gegenüber Patienten erwähnen): Sereina Urech, Martino Crivelli, Ali Peters.
Gegenüber Patienten ist JEDER im Team einfach "Therapeut" — keine Titel, keine Rollen, keine Hierarchien.

=== PATIENT FLOW (eine Frage auf einmal) ===
1. Was beschäftigt dich? (Symptome/Ziel)
2. Kurz empathisch reagieren — KEINE Diagnose stellen, keine medizinischen Ratschläge geben
3. Kurz empathisch, dann fragen ob Arzt/Bildgebung vorhanden — einladend, nicht als Hürde:
   Formulierung: "Warst du schon beim Arzt und hast eine Verordnung oder Bildgebung (Röntgen/MRI)? Falls du direkt zu uns kommen möchtest geht das auch — du kannst als Selbstzahler starten oder die Verordnung nachreichen."
   KOSTENTRÄGER — immer unterscheiden:
   - Unfall (Sturz, Verdrehung, Kollision, plötzliches Ereignis) → SUVA oder private Unfallversicherung — kein Arzt nötig, Therapeut meldet direkt
   - Krankheit (schleichend, chronisch, ohne klares Ereignis) → Grundversicherung, Arztverordnung nötig
   - Wenn Patient "verdreht", "gestürzt", "umgeknickt", "Zusammenstoss" erwähnt → sofort als Unfall einordnen, SUVA erwähnen
   - NIE pauschal "Krankenkasse" sagen wenn es ein Unfall sein könnte
4. Standort subtil klären: frage nach Wohnort/Arbeitsort/ÖV — NIE direkt nach Kineo-Standort fragen
5. Erst NACH Standortklärung: passende Therapeuten nennen — ALLE passenden, KEINEN bevorzugen
4. Sobald klar: "Ich lade die Termine..." + [LOAD_SLOTS:Standortname]
5. Patient wählt Slot → SMS-Auth → Buchung
Bei HYROX-Fragen: [LOAD_HYROX]
Verweise NIEMALS auf externe Buttons. Frage NIE nach Patientendaten im Chat.

=== STANDORTE & THERAPEUTEN ===
SEEFELD (Seefeldstrasse 83, 8008): Andrina Kümin, Joëlle Ramseier, Helen Schwank, Noah Stierli, Sereina Urech, Meike Vogel
→ Runners Lab: CURREX 5D Laufanalyse (einzigartig CH), Laufsport-Shop, kineo-runnerslab.ch

WIPKINGEN (Röschibachstrasse 79, 8037): Eva Danko, Raphael Hahner, Sonia Montero, Eve Schreurs, Barbara Victorio, Martino Crivelli, Max Boll
→ Athletiktraining, Laktatstufentest, Game Ready Kältetherapie

STAUFFACHER (Glasmalergasse 5, 8004): Andrina Kümin, Emma Leu, Carmen Weber
→ Functional Training, Personal Training

ESCHER WYSS (Hardturmstrasse 18, 8005): Clara Benning, Annika Heinrich, Andreas Niggl, Lucretia
→ HYROX Hub, Nina Schulte (Sportwissenschaftlerin): nach Physio 30 Min gratis Training = effektiv 60 Min
→ Kein Schwangerschaftslager, kein Sensopro. Sportkleidung mitbringen. kineo-hyrox.ch

THALWIL (Gotthardstrasse 52, 8800): Emma Leu, Andreas Niggl, Hanna Raffeiner, Joëlle Ramseier, Theresa Bitterlich
→ Kineo Fitness: eGym, Sensopro, fle-xx, Kurse (Yoga, Pilates, Aerial Yoga, TRX, HIIT). kineo-fitness.ch

ZOLLIKON (Zumikerstrasse 18, 8702): Helen Schwank, Meike Vogel
→ Yoga & Pilates von Physiotherapeuten geleitet. Kein Sensopro.

=== SENSOPRO ===
Vorhanden: Seefeld, Wipkingen, Stauffacher, Thalwil. NICHT bei Escher Wyss und Zollikon.

=== SCHWANGERSCHAFT ===
Schwangerschaftslager: alle Standorte AUSSER Escher Wyss.

=== LEISTUNGEN (Auswahl) ===
Physiotherapie, Manuelle Therapie, Dry Needling, Stosswellentherapie, CMD Kiefertherapie, Neurologie/Bobath, Lymphdrainage, Laufanalyse CURREX 5D, Runners Physio, Sportphysio, Beckenboden, Schwangerschaft, Faszientherapie, Digitales Heimprogramm, Medizinische Massage, Kältetherapie Game Ready, Laktatstufentest (Wipkingen), Athletiktraining, Sensopro

=== HYROX KURSE & PREISE ===
Coach: Juan Carlos Canache Fernandez
Montag: 06:30 Hyrox, 18:00 Full Body Strength
Dienstag: 06:30 Full Body Strength, 19:15 Hyrox
Mittwoch: 19:15 BOOTY L.B. Strength
Donnerstag: 09:50 Full Body Strength, 19:15 Hyrox
Freitag: 11:00 Hyrox
Personal Training: CHF 160/Std, 10er Pack CHF 1'400, Silver 8x/Mt CHF 1'040, Gold 12x/Mt CHF 1'440
Gruppenklassen-Preise: auf kineo-hyrox.ch/pricing
Buchung: sportsnow.ch/go/kineo-hyrox/classes (SportNow, nicht OneDoc)

=== KINEO FITNESS KURSE ===
5er Karte CHF 160, Schwangerschaft 10er CHF 280 (8 Mt). Schnupperkurs gratis. Abo jederzeit kündbar.
info@kineo-fitness.ch | 044 589 68 82

=== PREISE PHYSIOTHERAPIE ===
CHF 80 pro 30 Min. Mit ärztlicher Verordnung: Grundversicherung (10% Selbstbehalt nach Franchise). Klassische Massage: Zusatzversicherung (ausser Helsanagroup). Escher Wyss: 60 Min zum Preis von 30 Min (inkl. Nina-Training).

=== SPRACHEN ===
Deutsch (alle), Englisch (alle), Französisch (Wipkingen/Seefeld), Spanisch (Wipkingen/Thalwil), Italienisch (Wipkingen/Thalwil), Ungarisch (Wipkingen)

=== KONTAKT ===
Web: kineo-physiotherapie.ch (DE/EN/FR/IT) | info@kineo.swiss | 044 512 91 47
Runners Lab: kineo-runnerslab.ch | info@kineo-running.ch | 044 500 87 22
Instagram: @kineo.physiotherapie

=== TEAM-MODUS: STELLENPLAN & RESSOURCEN ===
Nur nach korrektem Team-PIN zugänglich. Stand: Dezember 2025/Frühjahr 2026.

ZIMMERPRINZIP: 1 Person = 1 Zimmer pro Tag. FD = Frühdienst, SD = Spätdienst. Max. 2 Personen/Tag/Zimmer (FD+SD).

PLANUNGSREGELN — ZWINGEND:
1. NUR NEUE MITARBEITER PLANEN — bestehende Therapeuten nie umplanen
2. IST-ZUSTAND — immer vom aktuellen Stand ausgehen, nie von Ideen-Szenarien
3. 1 TAG = 1 SCHICHT — auch wenn FD+SD frei: neue Person wählt nur eine davon pro Tag
4. FD/SD AUSGEWOGEN — bei mehreren Tagen: 50/50 FD/SD, z.B. 80% = 2 FD + 2 SD
5. MAXIMALE % KORREKT — max Stellen = Anzahl freie TAGE (nicht Slots)
6. IMMER KONKRET — Tagesplan angeben: Mo FD Z4 · Di SD Z2 · etc.
7. STELLENGRÖSSE — standardmässig 100% und 80% Stellen zeigen.
   Erst auf Rückfrage ("auch kleinere Stellen?") die 60% und 40% Optionen anzeigen.
   Beispiel-Einstieg: "Für eine 100%- oder 80%-Stelle passen folgende Kombinationen…"
6. STELLENGRÖSSE — immer zuerst 100% und 80% Optionen zeigen. Dann fragen: "Soll ich auch kleinere Pensen (60%/40%) anzeigen?" Geteilte Stellen nur auf explizite Anfrage.

=== THALWIL — IST-ZUSTAND (Mai–Aug) ===
Z1: Hanna Mo–Fr | Z2: Valerio Mo/Mi/Fr + Joëlle Di/Do | Z3: Emma Mo/Mi/Fr | Z4: FREI alle Tage
Freie Slots: Mo/Mi/Fr nur Z4 — Di/Do: Z3+Z4
Max neu: 140% (100%+40% ODER 80%+60%)
Kombination A: 100% Z4 Mo–Fr + 40% Z3 Di+Do
Kombination B: 80% Z4 Mo–Do + 60% Z3 Di+Do+Fr Z4

=== THALWIL — AB SEPTEMBER ===
Valerio 40% (nur Di Z2+Do Z2) · Joëlle weg nach Stauffacher/Seefeld
Z1: Hanna · Z2: FREI Mo/Mi/Fr + Valerio Di/Do · Z3: Emma Mo/Mi/Fr · Z4: FREI alle Tage
Max neu: 200% — Kombi A: 2×100% | B: 100%+60%+40% | C: 100%+80%+20% | D: 80%+80%+20%

=== ESCHER WYSS — IST-ZUSTAND (Mai–Aug) ===
Z1: FREI Mo + Clara Di–Fr · Z2: FREI · Z3: Lucretia Mo–Do
Max neu: 140% — Kombi A: 100% Z2 Mo–Fr + 40% Z1 Mo+Fr | B: 80% Z2 Mo–Do + 40% Z1 Mo+Fr

=== ESCHER WYSS — AB SEPTEMBER ===
Barbara Mo Z2 + Mi Z2 · Andrina Mo Z1 + Do Z2 · Lucretia bleibt Z3 Escher Wyss
Z1: Andrina Mo · Z2: Barbara Mo/Mi + Andrina Do · Z3: Lucretia Mo–Do
Freie Slots: nur Di Z2 + Fr Z2/Z3
Max neu: 60% — Kombi A: 40%+20% (Di Z2+Fr Z2/Z3)

=== WIPKINGEN — IST-ZUSTAND ===
Z1: Eva FD · Max SD | Z2: Raphael FD · Raphael SD | Z3: Barbara FD · Sonia SD (wechselnd) | Z4: Eve FD
Fr: kein SD ausser Sonia Z3
Freie Slots: Fr FD kein SD · Do FD kein freier Slot
Max neu: 140% — Kombi A: 80%+60% | B: 80%+60% | C: 80%+40% | D: 80%+20%

=== WIPKINGEN — AB SEPTEMBER ===
Barbara Di SD Z3 + Do SD Z3 · Meike Mo FD Z3 + Mi FD Z3 · Eve weg → Z4 komplett frei · Max bleibt Wipkingen
Z4 frei alle Tage · Z3 komplett belegt
Max neu: 140% → 100%+80% | 80%+80% | 100%+60% | 80%+60%
Maximum (260%): 100% Z4 + 80% Z2 (Mo FD/Di FD/Mi SD/Do SD) + 80% Z4 SD Tage

=== STAUFFACHER — IST-ZUSTAND ===
Z1: Andrina Mo+Do · Z2: Emma Di+Do · Carmen Fr
Freie Slots: Mo/Mi/Fr Z2 · Di/Mi Z1
Max neu: 100% — Kombi A: 60%+40% | B: 2×40%

=== STAUFFACHER — AB SEPTEMBER ===
Joëlle Di Z1 + Fr Z2 · Andrina Mo+Do weg nach Escher Wyss
Z1: Joëlle Di · Z2: Emma Di/Do + Joëlle Fr
Freie Slots: Mo Z1/Z2 · Mi Z1/Z2 · Do Z1
Max neu: 100% — Kombi A: 60% Mo/Mi/Do + 40% Mo/Mi | Kombi E: 60% allein Mo/Mi/Do

=== SEEFELD — IST-ZUSTAND ===
Z1: Sereina Mo/Di/Do + Helen Mi + Andrina Fr
Z2: Meike Mo/Mi + Helen Di + FREI Do
Z3: Joëlle Mo + Andrina Di/Mi + FREI Do + Joëlle Fr
Z4: Noah Mo/Di/Fr + FREI Mi/Do
Freie Slots: Di Z2 · Do Z2/Z3/Z4 — Mo/Mi/Fr komplett voll
Max neu: 40% (Di Z2 + Do Z2/Z3/Z4)

=== SEEFELD — AB SEPTEMBER ===
Joëlle Mo Z3+Mi Z4+Do Z3 neu · Andrina Di/Mi/Fr bleibt
Z2: FREI Mo/Mi → neue Stelle möglich
Variante B: Meike Di+Do statt Mo+Mi → Mo Z2+Mi Z2 frei
Max neu: 40% (Di Z2+Do Z4) oder 60% Variante B (Mo Z2+Mi Z2+Do Z4)

=== ZOLLIKON — IST-ZUSTAND (offiziell, Meike + Helen fix) ===
Z1: Meike Di+Do · Z2: Helen Di+Do · Z3: FREI alle Tage
Mo/Mi/Fr: Z1+Z2+Z3 alle frei · Di/Do: nur Z3 frei
Max neu für 100%/80% (Standard):
- 100%: P1 Z3 Mo–Fr (5 Tage, 2FD+3SD oder 3FD+2SD)
- 80%: P1 Z3 Mo–Do (4 Tage, 2FD+2SD)
Auf Rückfrage auch kleinere Stellen zeigen:
- 60%: Mo/Mi/Fr Z1 oder Z2 (wenn Meike/Helen nicht da)
- Max total mit kleinen Stellen: 100%+60%+60% = 220%

=== ZOLLIKON — AB SEPTEMBER (VARIANTE A) ===
Meike Di+Do weg → Helen bleibt Di+Do Z2
Z1: FREI · Z2: Helen Di+Do · Z3: FREI
Max neu: 260% — 3 Personen möglich Mo–Fr

=== ZOLLIKON — AB SEPTEMBER (VARIANTE B) ===
Meike Mo+Mi statt Di+Do · Helen Di+Do unverändert
Z1: Meike Mo+Mi · Z2: Helen Di+Do · Z3: frei
Kernkombination: Meike 40%+Helen 40%+P3 80% Z3 Mo–Do = 160%
Maximum: +P4 60% = 220% · +P4+P5 = 260%

=== SZENARIEN AB SEPTEMBER (INTERN / IDEENSPIEL) ===
WICHTIG: Diese Wechsel sind NICHT beschlossen und NICHT kommuniziert — nur ein internes Gedankenspiel.
NIE proaktiv erwähnen. Nur auf explizite Rückfrage einbeziehen ("Was wäre wenn Meike wechselt?" etc.).
Wenn gefragt: klar kommunizieren dass es sich um ein Szenario handelt, nicht um beschlossene Sache.

Valerio: Thalwil 100% → 40% (nur Di Z2+Do Z2)
Joëlle: Thalwil Di+Do → Stauffacher Di Z1+Fr Z2 · Seefeld Mo Z3+Mi Z4+Do Z3
Barbara: Mo+Mi Escher Wyss Z2 · Di SD+Do SD Wipkingen Z3
Meike: Di+Do Zollikon weg → Mo FD Z3+Mi FD Z3 Wipkingen
Andrina: Mo+Do Stauffacher weg → Mo Z1+Do Z2 Escher Wyss

=== MAXIMALE NEUE STELLEN ÜBERSICHT ===
Zollikon: max 220% neu (Variante A) / 260% (mit Variante B)
Thalwil: max 200% neu (ab September)
Escher Wyss: max 140% (Mai–Aug) / 60% (ab September)
Wipkingen: max 140% neu
Stauffacher: max 100% neu
Seefeld: max 40% neu (nur 1 Stelle)

WICHTIG: Alle Zimmer- und Stellendaten sind interne Informationen — nur nach PIN-Verifikation zeigen.

=== SPEZIALISIERUNGEN (Ja + Basic = empfehlenswert, aus offizieller Liste) ===
WICHTIG: Nur Therapeuten empfehlen die auf OneDoc buchbar sind.
Theresa Bitterlich scheidet Ende Mai aus — ab Juni NICHT mehr empfehlen.

Neurologie/Migräne/Kopfschmerz: Sereina Urech, Hanna Raffeiner, Carmen Weber, Andrina Kümin, Helen Schwank, Raphael Hahner, Theresa Bitterlich (bis Mai), Eva Danko, Andreas Niggl, Noah Stierli, Annika Heinrich, Sonia Montero, Clara Benning, Barbara Victorio, Meike Vogel, Emma Leu, Joëlle Ramseier
CMD/Kiefer: Sereina Urech, Hanna Raffeiner, Carmen Weber, Helen Schwank, Theresa Bitterlich (bis Mai), Eva Danko, Andreas Niggl, Noah Stierli, Barbara Victorio, Meike Vogel
Schwindel: Sereina Urech, Hanna Raffeiner, Carmen Weber, Andrina Kümin, Helen Schwank, Raphael Hahner, Andreas Niggl, Meike Vogel
Runners Physio: Sereina Urech, Hanna Raffeiner, Carmen Weber, Andrina Kümin, Helen Schwank, Raphael Hahner, Theresa Bitterlich (bis Mai), Eva Danko, Andreas Niggl, Noah Stierli
5D Laufanalyse (Runners Lab Seefeld): Sereina Urech, Andrina Kümin, Noah Stierli
Beckenboden: Helen Schwank, Theresa Bitterlich (bis Mai), Eva Danko (inkl. Biofeedback/Vaginalsonden)
Bobath: Carmen Weber
Sport: Hanna Raffeiner, Andrina Kümin, Raphael Hahner, Andreas Niggl, Noah Stierli, Annika Heinrich, Sonia Montero, Clara Benning, Barbara Victorio
Athletiktraining: Raphael Hahner, Eva Danko, Andreas Niggl, Clara Benning
Dry Needling: Sereina Urech, Hanna Raffeiner, Andrina Kümin, Raphael Hahner, Theresa Bitterlich (bis Mai), Eva Danko, Andreas Niggl, Annika Heinrich, Sonia Montero, Barbara Victorio
Lymphdrainage: Hanna Raffeiner, Helen Schwank, Raphael Hahner, Theresa Bitterlich (bis Mai), Meike Vogel, Emma Leu
Stosswelle: alle Therapeuten
Functional Training: Andrina Kümin, Raphael Hahner, Andreas Niggl, Annika Heinrich, Sonia Montero, Barbara Victorio
Yoga: Helen Schwank, Barbara Victorio, Meike Vogel
Pilates: Barbara Victorio
Re-Analyse (Laufanalyse): Sereina Urech, Andrina Kümin, Noah Stierli
Schwangerschaftsliege: alle Therapeuten an allen Standorten ausser Escher Wyss (kein Lager)
USZ CMCS Projekt: Hanna Raffeiner (Thalwil), Martino Crivelli (Wipkingen)

BEI EMPFEHLUNGEN:
- ALLE passenden nennen — von ALLEN Standorten — nie einzelne bevorzugen
- Standort erst danach klären
- Keine Hierarchie — alphabetisch oder rotierend
- Sereina nie als Leiterin/Managerin erwähnen — sie ist Therapeutin wie alle anderen

=== VERHALTEN ===
- Empfehle IMMER alle passenden Therapeuten — von allen Standorten — nie einen einzelnen bevorzugen
- Sereina NIE als Managerin oder besonders hervorheben — sie ist Therapeutin wie alle anderen
- Eine Frage auf einmal. Warm, direkt, empathisch.
- Bei Fotos/Bildern: auf Website oder Instagram verweisen
- Erfinde keine Verfügbarkeiten — nur echte Slots via [LOAD_SLOTS:X]
- Team-PIN NIEMALS im Chat anzeigen
"""


class Message(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    messages: List[Message]

class ExportRequest(BaseModel):
    messages: List[Message]
    scope: str = "all"


TEAM_PIN = os.environ.get("TEAM_PIN", "1234")

@app.post("/api/chat")
async def chat(req: ChatRequest):
    if not req.messages:
        raise HTTPException(400, "Keine Nachrichten")
    try:
        # Inject team PIN into system prompt dynamically
        system = SYSTEM_PROMPT.replace("${TEAM_PIN}", TEAM_PIN)
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=600,
            system=system,
            messages=[{"role": m.role, "content": m.content} for m in req.messages],
        )
        reply = response.content[0].text

        # Detect [LOAD_SLOTS:Standort] marker
        import re
        load_slots = None
        slot_match = re.search(r'\[LOAD_SLOTS:([^\]]+)\]', reply)
        if slot_match:
            load_slots = slot_match.group(1).strip()
            reply = re.sub(r'\[LOAD_SLOTS:[^\]]+\]', '', reply).strip()

        # Detect [LOAD_HYROX] marker
        load_hyrox = False
        if '[LOAD_HYROX]' in reply:
            load_hyrox = True
            reply = reply.replace('[LOAD_HYROX]', '').strip()

        # Signal team unlock to frontend
        team_unlocked = f"[TEAM_UNLOCKED:{TEAM_PIN}]" in " ".join(m.content for m in req.messages)
        return {"reply": reply, "team_mode": team_unlocked, "load_slots": load_slots, "load_hyrox": load_hyrox}
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/export-word")
async def export_word(req: ExportRequest):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from datetime import datetime

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        title = doc.add_heading('Kineo Stellenplan-Assistent', 0)
        title.runs[0].font.color.rgb = RGBColor(0x2C, 0x2C, 0x2A)
        title.runs[0].font.size = Pt(18)

        date_p = doc.add_paragraph(f'Exportiert am {datetime.now().strftime("%d.%m.%Y %H:%M")} Uhr')
        date_p.runs[0].font.size = Pt(10)
        date_p.runs[0].font.color.rgb = RGBColor(0x73, 0x72, 0x6C)
        doc.add_paragraph()

        for msg in req.messages:
            is_user = msg.role == "user"

            label = doc.add_paragraph()
            label_run = label.add_run("Frage" if is_user else "Antwort")
            label_run.bold = True
            label_run.font.size = Pt(9)
            label_run.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5) if is_user else RGBColor(0x73, 0x72, 0x6C)
            label.paragraph_format.space_after = Pt(2)

            content_p = doc.add_paragraph()
            content_p.paragraph_format.left_indent = Inches(0.2)
            content_p.paragraph_format.space_after = Pt(10)

            pPr = content_p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'E6F1FB' if is_user else 'F7F6F3')
            pPr.append(shd)

            run = content_p.add_run(msg.content)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5) if is_user else RGBColor(0x2C, 0x2C, 0x2A)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=kineo-stellenplan.docx"}
        )
    except ImportError:
        raise HTTPException(500, "python-docx nicht installiert")
    except Exception as e:
        raise HTTPException(500, str(e))




# Stellenplan-DOCX direkt als Download ausliefern
from fastapi.responses import FileResponse as FR
import pathlib

@app.get("/kineo_stellenplanung.docx")
async def download_stellenplan():
    path = pathlib.Path("kineo_stellenplanung.docx")
    if not path.exists():
        raise HTTPException(404, "Stellenplan nicht gefunden")
    return FR(
        path=str(path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="Kineo_Stellenplanung.docx"
    )


# ── THERAPEUTEN-PROFIL ENDPOINT ────────────────────────────────────────────────
import httpx, asyncio
from fastapi.responses import JSONResponse

THERAPEUTEN_URLS = {
    "Helen Schwank":    "https://www.kineo-physiotherapie.ch/en-ch/team/helen-schwank",
    "Sereina Urech":    "https://www.kineo-physiotherapie.ch/en-ch/team/sereina-urech",
    "Noah Stierli":     "https://www.kineo-physiotherapie.ch/en-ch/team/noah-stierli",
    "Joëlle Ramseier":  "https://www.kineo-physiotherapie.ch/en-ch/team/joelle-ramseier",
    "Andrina Kümin":    "https://www.kineo-physiotherapie.ch/en-ch/team/andrina-kumin",
    "Meike Vogel":      "https://www.kineo-physiotherapie.ch/en-ch/team/meike-vogel",
    "Felica Kossendey": "https://www.kineo-physiotherapie.ch/en-ch/team/felica-kossendey",
    "Emma":             "https://www.kineo-physiotherapie.ch/en-ch/team/emma",
    "Hanna":            "https://www.kineo-physiotherapie.ch/en-ch/team/hanna",
    "Barbara":          "https://www.kineo-physiotherapie.ch/en-ch/team/barbara",
    "Valerio":          "https://www.kineo-physiotherapie.ch/en-ch/team/valerio-sasso",
}

# Google Places API key (optional – set as env var GOOGLE_PLACES_KEY)
GOOGLE_KEY = os.environ.get("GOOGLE_PLACES_KEY", "")

# Google Place IDs for each Kineo location
PLACE_IDS = {
    "Seefeld":      "ChIJe4tGmjQKkEcR9K-FV2r8Guc",
    "Wipkingen":    "ChIJmyJeEgcKkEcRoFkZpFmJqOI",
    "Stauffacher":  "ChIJcR6HijkKkEcRFz5xhXrFpS4",
    "Thalwil":      "ChIJV8C3Y6wPkEcRBCX1dHdFpek",
    "Escher Wyss":  "ChIJT0KFUQcKkEcR4DLriJGf4dk",
    "Zollikon":     "ChIJk2WS1OUPkEcRHwJw3F3cLCo",
}

async def fetch_therapist(session: httpx.AsyncClient, name: str, url: str) -> dict:
    """Fetch leistungen from Kineo homepage for one therapist."""
    try:
        r = await session.get(url, timeout=6,
            headers={"User-Agent": "Mozilla/5.0 (compatible; KineoBot/1.0)"})
        import re
        text = re.sub(r'<[^>]+>', ' ', r.text)
        text = re.sub(r'\s+', ' ', text)

        # Extract leistungen — they appear as linked list items on the page
        leistungen = re.findall(
            r'(?:General Physiotherapy|Neurology|CMD Jaw Therapy|Shockwave|Dry Needling|'
            r'Runners Physio|Sports Physio|Lymph|Vertigo|Massage|Laufanalyse|'
            r'Manual Therapy|Pilates|Kinesio)',
            text, re.IGNORECASE
        )
        # Deduplicate preserving order
        seen = set()
        leistungen = [x for x in leistungen if not (x.lower() in seen or seen.add(x.lower()))]
        return {"name": name, "leistungen": leistungen, "url": url, "source": "homepage"}
    except Exception as e:
        return {"name": name, "leistungen": [], "url": url, "source": "error", "error": str(e)}


async def fetch_google_reviews(session: httpx.AsyncClient, standort: str, place_id: str) -> dict:
    """Fetch Google rating for one Kineo location."""
    if not GOOGLE_KEY:
        return {"standort": standort, "note": "kein Google API Key konfiguriert"}
    try:
        url = (
            f"https://maps.googleapis.com/maps/api/place/details/json"
            f"?place_id={place_id}&fields=rating,user_ratings_total,reviews"
            f"&language=de&key={GOOGLE_KEY}"
        )
        r = await session.get(url, timeout=6)
        data = r.json().get("result", {})
        reviews = data.get("reviews", [])[:5]  # Top 5
        return {
            "standort": standort,
            "rating": data.get("rating"),
            "total": data.get("user_ratings_total"),
            "top_reviews": [
                {"autor": rv.get("author_name"), "sterne": rv.get("rating"),
                 "text": rv.get("text","")[:200]}
                for rv in reviews
            ]
        }
    except Exception as e:
        return {"standort": standort, "error": str(e)}


@app.get("/api/therapeuten")
async def get_therapeuten():
    """Live-Abfrage: Leistungen von Kineo-Homepage + Google-Bewertungen."""
    async with httpx.AsyncClient(follow_redirects=True) as session:
        # Therapeuten-Profile parallel abrufen
        tasks_t = [fetch_therapist(session, n, u) for n, u in THERAPEUTEN_URLS.items()]
        tasks_g = [fetch_google_reviews(session, s, p) for s, p in PLACE_IDS.items()]
        results = await asyncio.gather(*tasks_t, *tasks_g, return_exceptions=True)

    n = len(THERAPEUTEN_URLS)
    therapeuten = [r for r in results[:n] if isinstance(r, dict)]
    bewertungen  = [r for r in results[n:] if isinstance(r, dict)]

    return JSONResponse({
        "therapeuten": therapeuten,
        "standort_bewertungen": bewertungen,
        "hinweis": "Leistungen live von kineo-physiotherapie.ch. Google nur mit GOOGLE_PLACES_KEY."
    })


# ── SMS VERIFIKATION & IDENTIFIKATION ────────────────────────────────────────
import secrets, time, hashlib
from pydantic import BaseModel

# Env vars: SMS_GW_USER, SMS_GW_PASS (nie im Code!)
TWILIO_SID   = os.environ.get("TWILIO_ACCOUNT_SID", "")
TWILIO_TOKEN = os.environ.get("TWILIO_AUTH_TOKEN", "")
TWILIO_FROM  = os.environ.get("TWILIO_FROM", "")
SMS_SENDER   = os.environ.get("SMS_SENDER", "Kineo")

# In-Memory OTP Store: {session_id: {code, phone, expires, verified, name, dob}}
# Für Produktion: Redis verwenden
otp_store: dict = {}

def generate_otp() -> str:
    return str(secrets.randbelow(900000) + 100000)  # 6-stellig

async def send_sms(phone: str, message: str) -> bool:
    """SMS via Twilio."""
    if not TWILIO_SID or not TWILIO_TOKEN or not TWILIO_FROM:
        print("Twilio credentials nicht gesetzt")
        return False

    dest = phone.strip().replace(" ", "").replace("-", "")
    if dest.startswith("0") and not dest.startswith("00"):
        dest = "+41" + dest[1:]
    if not dest.startswith("+"):
        dest = "+" + dest

    print(f"SMS → dest:{dest} from:{TWILIO_FROM}")

    try:
        async with httpx.AsyncClient(timeout=12) as session:
            r = await session.post(
                f"https://api.twilio.com/2010-04-01/Accounts/{TWILIO_SID}/Messages.json",
                auth=(TWILIO_SID, TWILIO_TOKEN),
                data={
                    "To":   dest,
                    "From": TWILIO_FROM,
                    "Body": message,
                },
            )
            print(f"Twilio → {r.status_code}: {r.text[:300]}")
            if not r.text.strip():
                print("Twilio: leere Antwort — Auth Token prüfen")
                return False
            data = r.json()
            print(f"Twilio status={data.get('status')} error={data.get('message','')}")
            return r.status_code in (200, 201) and data.get("status") in ("queued", "sent", "delivered")
    except Exception as e:
        print(f"SMS Fehler: {e}")
        return False


class OtpRequest(BaseModel):
    phone: str
    name: str
    dob: str          # Format: DD.MM.YYYY

class OtpVerify(BaseModel):
    session_id: str
    code: str

class PatientMode(BaseModel):
    mode: str         # "patient" oder "staff"
    pin: str = ""     # Staff-PIN (optional)


@app.post("/api/auth/send-otp")
async def send_otp(req: OtpRequest):
    """Schritt 1: OTP per SMS senden."""
    # Telefonnummer normalisieren
    phone = req.phone.strip().replace(" ", "").replace("-", "")
    if phone.startswith("0") and not phone.startswith("00"):
        phone = "+41" + phone[1:]  # Schweizer Nummer
    if not phone.startswith("+"):
        phone = "+" + phone

    code = generate_otp()
    session_id = secrets.token_urlsafe(16)
    expires = time.time() + 300  # 5 Minuten

    otp_store[session_id] = {
        "code": code,
        "phone": phone,
        "name": req.name.strip(),
        "dob": req.dob.strip(),
        "expires": expires,
        "verified": False,
        "attempts": 0,
    }
    # Auch in DB speichern (persistent, optional)
    try:
        await db_save_otp(session_id, phone, req.name.strip(), req.dob.strip(), code, expires)
    except Exception as e:
        print(f"DB save OTP fehler (nicht kritisch): {e}")

    # Alte Sessions aufräumen
    now = time.time()
    expired = [k for k, v in otp_store.items() if v["expires"] < now]
    for k in expired:
        del otp_store[k]

    msg = f"Ihr Kineo-Code: {code}. Gültig 5 Minuten. Nicht weitergeben."
    sent = await send_sms(phone, msg)

    if not sent and TWILIO_SID:
        return JSONResponse({"error": "SMS konnte nicht gesendet werden"}, status_code=500)

    return JSONResponse({
        "session_id": session_id,
        "phone_masked": phone[:4] + "****" + phone[-3:],
        "expires_in": 300,
        "sms_sent": sent,
        **({"dev_code": code} if (not TWILIO_SID or not sent) else {}),
    })


@app.post("/api/auth/verify-otp")
async def verify_otp(req: OtpVerify):
    """Schritt 2: OTP verifizieren."""
    session = otp_store.get(req.session_id)
    if not session:
        return JSONResponse({"error": "Session nicht gefunden oder abgelaufen"}, status_code=400)
    if time.time() > session["expires"]:
        del otp_store[req.session_id]
        return JSONResponse({"error": "Code abgelaufen — bitte neu anfordern"}, status_code=400)
    if session["attempts"] >= 3:
        return JSONResponse({"error": "Zu viele Versuche — bitte neu anfordern"}, status_code=400)

    session["attempts"] += 1
    if req.code.strip() != session["code"]:
        remaining = 3 - session["attempts"]
        return JSONResponse({"error": f"Falscher Code. Noch {remaining} Versuch(e)."}, status_code=400)

    # Verifiziert!
    session["verified"] = True
    auth_token = secrets.token_urlsafe(24)
    session["auth_token"] = auth_token
    session["verified_at"] = time.time()

    # Auch in DB verifizieren (optional)
    try:
        db_result = await db_verify_otp(req.session_id, req.code.strip())
        if db_result and "auth_token" in db_result:
            auth_token = db_result["auth_token"]
            session["auth_token"] = auth_token
    except Exception as e:
        print(f"DB verify fehler (nicht kritisch): {e}")

    return JSONResponse({
        "success": True,
        "auth_token": auth_token,
        "patient": {
            "name": session["name"],
            "dob": session["dob"],
            "phone": session["phone"],
        }
    })


# ── ONEDOC BUCHUNG ENDPOINT ───────────────────────────────────────────────────

class BookingRequest(BaseModel):
    session_id: str
    auth_token: str
    # Slot-Details (aus fetch_onedoc_slots)
    entity_id: str
    prof_id: str
    calendar_id: str
    appointment_type_id: str
    datetime_iso: str       # z.B. "2026-04-15T09:00:00+02:00"
    therapeut: str
    standort: str
    # Patientendaten (aus verifizierter Session)
    vorname: str = ""
    nachname: str = ""
    geburtsdatum: str = ""  # DD.MM.YYYY
    telefon: str = ""
    email: str = ""
    kommentar: str = ""


@app.post("/api/buchen")
async def buchen(req: BookingRequest):
    """Bucht einen Termin direkt in OneDoc — kein Redirect, keine externe Seite."""

    # 1. Auth prüfen
    session = otp_store.get(req.session_id)
    if not session or not session.get("verified"):
        return JSONResponse({"error": "Nicht authentifiziert"}, status_code=401)
    if session.get("auth_token") != req.auth_token:
        return JSONResponse({"error": "Ungültiges Auth-Token"}, status_code=401)
    if time.time() > session["verified_at"] + 3600:
        return JSONResponse({"error": "Session abgelaufen"}, status_code=401)

    # 2. Patientendaten aus Session wenn nicht übergeben
    name_parts = session["name"].split(" ", 1)
    vorname   = req.vorname  or (name_parts[0] if name_parts else "")
    nachname  = req.nachname or (name_parts[1] if len(name_parts) > 1 else "")
    geburtsdatum = req.geburtsdatum or session.get("dob", "")
    telefon   = req.telefon  or session.get("phone", "")

    # Geburtsdatum: DD.MM.YYYY → YYYY-MM-DD
    dob_iso = geburtsdatum
    if geburtsdatum and "." in geburtsdatum:
        parts = geburtsdatum.split(".")
        if len(parts) == 3:
            dob_iso = f"{parts[2]}-{parts[1].zfill(2)}-{parts[0].zfill(2)}"

    # 3. Buchungs-Payload — gleiche Struktur wie OneDoc-Website POST
    payload = {
        "appointmentTypeId": int(req.appointment_type_id),
        "calendarId":        int(req.calendar_id),
        "professionalId":    int(req.prof_id),
        "entityId":          int(req.entity_id),
        "startDateTime":     req.datetime_iso,
        "patient": {
            "firstName":   vorname,
            "lastName":    nachname,
            "birthDate":   dob_iso,
            "phoneNumber": telefon,
            "email":       req.email or "",
        },
        "acceptedClientType": "established",
        "comment": req.kommentar or "",
        "source": "web-patient",
    }

    headers = {
        **ONEDOC_HEADERS,
        "Content-Type": "application/json",
        "Origin": "https://www.onedoc.ch",
    }

    # Route booking through Hetzner proxy (Render can't reach OneDoc directly)
    async with httpx.AsyncClient(follow_redirects=True) as session_http:
        try:
            proxy_headers = {}
            if SLOTS_API_KEY:
                proxy_headers["X-API-Key"] = SLOTS_API_KEY
            r = await session_http.post(
                f"{SLOTS_PROXY}/buchen",
                json=payload,
                headers=proxy_headers,
                timeout=25,
            )

            if r.status_code in (200, 201):
                data = r.json()
                appointment_id = (
                    data.get("data", {}).get("id") or
                    data.get("id") or
                    data.get("appointmentId") or
                    "unbekannt"
                )
                # In Session + DB speichern
                sess = otp_store.get(req.session_id, {})
                if "appointments" not in sess:
                    sess["appointments"] = []
                    sess["appointments_detail"] = []
                sess["appointments"].append(str(appointment_id))
                sess["appointments_detail"].append({
                    "id": str(appointment_id),
                    "therapeut": req.therapeut,
                    "standort": req.standort,
                    "datum": req.datetime_iso[:10],
                    "zeit": req.datetime_iso[11:16],
                    "gebucht_am": time.strftime("%Y-%m-%d %H:%M"),
                })
                try:
                    await db_save_appointment(req.session_id, {
                        "patient_name": f"{vorname} {nachname}",
                        "patient_phone": telefon,
                        "patient_dob": dob_iso,
                        "appointment_id": appointment_id,
                        "therapeut": req.therapeut,
                        "standort": req.standort,
                        "datum": req.datetime_iso[:10],
                        "zeit": req.datetime_iso[11:16],
                        "datetime_iso": req.datetime_iso,
                        "entity_id": req.entity_id,
                        "prof_id": req.prof_id,
                        "calendar_id": req.calendar_id,
                    })
                except Exception as e:
                    print(f"DB appointment save fehler (nicht kritisch): {e}")

                return JSONResponse({
                    "success": True,
                    "appointment_id": appointment_id,
                    "therapeut": req.therapeut,
                    "standort": req.standort,
                    "datum": req.datetime_iso[:10],
                    "zeit": req.datetime_iso[11:16],
                    "patient": f"{vorname} {nachname}",
                    "message": f"Termin gebucht: {req.therapeut}, {req.datetime_iso[:10]} um {req.datetime_iso[11:16]} Uhr",
                })

            # OneDoc gibt manchmal 422 mit Fehlermeldung zurück
            error_body = {}
            try:
                error_body = r.json()
            except Exception:
                pass

            error_msg = (
                error_body.get("message") or
                error_body.get("error") or
                error_body.get("errors", [{}])[0].get("message", "") or
                f"HTTP {r.status_code}"
            )

            return JSONResponse({
                "error": error_msg,
                "status": r.status_code,
                "detail": error_body,
            }, status_code=400)

        except httpx.TimeoutException:
            return JSONResponse({"error": "Timeout — OneDoc nicht erreichbar"}, status_code=504)
        except Exception as e:
            return JSONResponse({"error": str(e)}, status_code=500)


# ── ONEDOC ABSAGE ENDPOINT ────────────────────────────────────────────────────

class CancelRequest(BaseModel):
    session_id: str
    auth_token: str
    appointment_id: str   # von OneDoc bei Buchung zurückgegeben
    grund: str = ""       # optional

@app.post("/api/absagen")
async def absagen(req: CancelRequest):
    """Termin bei OneDoc stornieren."""

    # Auth prüfen
    session = otp_store.get(req.session_id)
    if not session or not session.get("verified"):
        return JSONResponse({"error": "Nicht authentifiziert"}, status_code=401)
    if session.get("auth_token") != req.auth_token:
        return JSONResponse({"error": "Ungültiges Auth-Token"}, status_code=401)
    if time.time() > session["verified_at"] + 3600:
        return JSONResponse({"error": "Session abgelaufen"}, status_code=401)

    # Aus DB prüfen ob Termin dem Patient gehört (auch nach Browser-Neustart)
    phone = session.get("phone", "")
    db_owned = await db_cancel_appointment(req.appointment_id, req.session_id, phone)
    # Auch In-Memory check als Fallback
    in_memory_owned = req.appointment_id in session.get("appointments", [])
    if not db_owned and not in_memory_owned:
        return JSONResponse({
            "error": "Termin nicht gefunden — du kannst nur deine eigenen Termine absagen"
        }, status_code=403)

    async with httpx.AsyncClient(follow_redirects=True) as http:
        try:
            # Versuch 1: DELETE
            r = await http.delete(
                f"{ONEDOC_BASE}/api/v1/appointments/{req.appointment_id}",
                headers={**ONEDOC_HEADERS, "Content-Type": "application/json"},
                json={"cancellationReason": req.grund or "Abgesagt durch Patient"},
                timeout=15,
            )

            # Versuch 2: PATCH mit status=cancelled (falls DELETE nicht supported)
            if r.status_code == 405:
                r = await http.patch(
                    f"{ONEDOC_BASE}/api/v1/appointments/{req.appointment_id}",
                    headers={**ONEDOC_HEADERS, "Content-Type": "application/json"},
                    json={"status": "cancelled", "cancellationReason": req.grund or "Abgesagt durch Patient"},
                    timeout=15,
                )

            if r.status_code in (200, 201, 204):
                # Aus Session entfernen
                session["appointments"] = [
                    a for a in patient_appointments if a != req.appointment_id
                ]
                return JSONResponse({
                    "success": True,
                    "appointment_id": req.appointment_id,
                    "message": "Termin wurde erfolgreich abgesagt.",
                })

            error_body = {}
            try: error_body = r.json()
            except Exception: pass

            error_msg = (
                error_body.get("message") or
                error_body.get("error") or
                f"HTTP {r.status_code}"
            )
            return JSONResponse({"error": error_msg, "status": r.status_code}, status_code=400)

        except httpx.TimeoutException:
            return JSONResponse({"error": "Timeout — OneDoc nicht erreichbar"}, status_code=504)
        except Exception as e:
            return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/meine-termine")
async def meine_termine(session_id: str, token: str):
    """Gibt gebuchte Termine des verifizierten Patienten zurück."""
    session = otp_store.get(session_id)
    if not session or not session.get("verified") or session.get("auth_token") != token:
        return JSONResponse({"error": "Nicht authentifiziert"}, status_code=401)

    # Aus DB laden (auch nach Browser-Neustart verfügbar)
    phone = session.get("phone", "")
    db_termine = await db_get_appointments(session_id, phone)
    if db_termine:
        termine = [{
            "id": str(t["appointment_id"]),
            "therapeut": t["therapeut"],
            "standort": t["standort"],
            "datum": t["datum"],
            "zeit": t["zeit"],
            "gebucht_am": str(t["gebucht_am"])[:16] if t["gebucht_am"] else "",
        } for t in db_termine]
    else:
        termine = session.get("appointments_detail", [])
    return JSONResponse({
        "patient": session["name"],
        "termine": termine,
    })


@app.get("/api/auth/status")
async def auth_status(session_id: str, token: str):
    """Prüft ob eine Session noch gültig und verifiziert ist."""
    session = otp_store.get(session_id)
    if not session:
        return JSONResponse({"authenticated": False})
    if not session.get("verified"):
        return JSONResponse({"authenticated": False})
    if session.get("auth_token") != token:
        return JSONResponse({"authenticated": False})
    if time.time() > session["verified_at"] + 3600:  # 1h Session-Timeout
        return JSONResponse({"authenticated": False, "reason": "session_expired"})
    return JSONResponse({
        "authenticated": True,
        "patient": {
            "name": session["name"],
            "dob": session["dob"],
        }
    })


# ── ONEDOC LIVE-TERMINE ENDPOINT ──────────────────────────────────────────────
# OneDoc Pro API: freie Buchungsslots pro Standort live abrufen
# Env vars benötigt: ONEDOC_API_KEY (Pflicht), ONEDOC_PRACTICE_IDS (optional, JSON)

ONEDOC_KEY = os.environ.get("ONEDOC_API_KEY", "")

# OneDoc Practice-IDs für jeden Kineo-Standort
# Format: {"Seefeld": "PRACTICE_ID", ...}
# Kann als Env-Variable ONEDOC_PRACTICE_IDS (JSON-String) überschrieben werden
import json as _json
# Vollständige Therapeuten-Daten aus scraper_api.py
PRAXEN_DATA = [
    {"standort": "Seefeld", "entity_id": "50967", "therapeuten": [
        {"name": "Andrina Kümin",   "prof_id": "2907953", "calendar_id": "125442"},
        {"name": "Joëlle Ramseier", "prof_id": "2998987", "calendar_id": "128306"},
        {"name": "Helen Schwank",   "prof_id": "2907961", "calendar_id": "125445"},
        {"name": "Noah Stierli",    "prof_id": "2907964", "calendar_id": "125443"},
        {"name": "Sereina Urech",   "prof_id": "2907969", "calendar_id": "125444"},
        {"name": "Meike Vogel",     "prof_id": "2944269", "calendar_id": "126599"},
    ]},
    {"standort": "Wipkingen", "entity_id": "50970", "therapeuten": [
        {"name": "Eva Danko",        "prof_id": "2907150", "calendar_id": "125383"},
        {"name": "Raphael Hahner",   "prof_id": "2907166", "calendar_id": "125386"},
        {"name": "Sonia Montero",    "prof_id": "2907167", "calendar_id": "125387"},
        {"name": "Eve Schreurs",     "prof_id": "2999078", "calendar_id": "128319"},
        {"name": "Barbara Victorio", "prof_id": "2944281", "calendar_id": "126601"},
    ]},
    {"standort": "Stauffacher", "entity_id": "50968", "therapeuten": [
        {"name": "Andrina Kümin",   "prof_id": "2908008", "calendar_id": "125453"},
        {"name": "Emma Leu",        "prof_id": "2962676", "calendar_id": "127231"},
        {"name": "Carmen Weber",    "prof_id": "2911916", "calendar_id": "125641"},
    ]},
    {"standort": "Escher Wyss", "entity_id": "51318", "therapeuten": [
        {"name": "Clara Benning",   "prof_id": "2936879", "calendar_id": "126376"},
        {"name": "Annika Heinrich", "prof_id": "2936895", "calendar_id": "126377"},
        {"name": "Andreas Niggl",   "prof_id": "2936885", "calendar_id": "126378"},
    ]},
    {"standort": "Zollikon", "entity_id": "50971", "therapeuten": [
        {"name": "Helen Schwank", "prof_id": "2907110", "calendar_id": "125377"},
        {"name": "Meike Vogel",   "prof_id": "2944266", "calendar_id": "126598"},
    ]},
    {"standort": "Thalwil", "entity_id": "50969", "therapeuten": [
        {"name": "Theresa Bitterlich", "prof_id": "2907175", "calendar_id": "125393"},
        {"name": "Emma Leu",           "prof_id": "2962677", "calendar_id": "127230"},
        {"name": "Andreas Niggl",      "prof_id": "2907180", "calendar_id": "125439"},
        {"name": "Hanna Raffeiner",    "prof_id": "2907170", "calendar_id": "125395"},
        {"name": "Joëlle Ramseier",    "prof_id": "3056625", "calendar_id": "129911"},
    ]},
]

# Helper lookups
_entity_by_standort = {p["standort"]: p["entity_id"] for p in PRAXEN_DATA}
_therapeut_by_name  = {
    t["name"]: {**t, "standort": p["standort"], "entity_id": p["entity_id"]}
    for p in PRAXEN_DATA for t in p["therapeuten"]
}

_default_practice_ids = {p["standort"]: p["entity_id"] for p in PRAXEN_DATA}
try:
    ONEDOC_PRACTICE_IDS = _json.loads(os.environ.get("ONEDOC_PRACTICE_IDS", "{}"))
    # Merge mit defaults
    for k in _default_practice_ids:
        if k not in ONEDOC_PRACTICE_IDS:
            ONEDOC_PRACTICE_IDS[k] = _default_practice_ids[k]
except Exception:
    ONEDOC_PRACTICE_IDS = _default_practice_ids


ONEDOC_BASE = "https://www.onedoc.ch"
ONEDOC_HEADERS = {
    "User-Agent":       "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/26.3 Safari/605.1.15",
    "Accept":           "application/json, */*",
    "Accept-Language":  "de-DE,de;q=0.9",
    "Referer":          "https://www.onedoc.ch/de/",
    "X-Requested-With": "XMLHttpRequest",
    "X-API-Version":    "1.2.0",
    "X-OneDoc-App":     "web-patient",
}
STANDARD_APT_TYPE_ID = 176378  # Allg. Physiotherapie 30min


async def onedoc_get(session: httpx.AsyncClient, path: str, params: dict = None):
    try:
        r = await session.get(
            ONEDOC_BASE + path, params=params,
            headers=ONEDOC_HEADERS, timeout=30, follow_redirects=True,
        )
        return r.status_code, (r.json() if r.status_code == 200 else None)
    except Exception as e:
        print(f"OneDoc GET Fehler: {type(e).__name__}: {e}")
        return 0, None


async def get_apt_type_id(session: httpx.AsyncClient, entity_id: str, prof_id: str) -> str:
    for ct in ["established", "new"]:
        status, data = await onedoc_get(
            session, f"/api/entities/{entity_id}/bookable-appointment-types",
            {"professionalId": prof_id, "acceptedClientType": ct},
        )
        if status == 200 and data:
            typen = data.get("data", [])
            for t in typen:
                if t.get("appointmentTypeId") == STANDARD_APT_TYPE_ID:
                    return str(STANDARD_APT_TYPE_ID)
            if typen:
                return str(typen[0]["appointmentTypeId"])
    return str(STANDARD_APT_TYPE_ID)


async def fetch_onedoc_slots(
    session: httpx.AsyncClient,
    standort: str,
    practice_id: str,
    days_ahead: int = 14,
    prof_id: str = None,
) -> dict:
    """Freie Slots via echter OneDoc-API (gleicher Endpunkt wie scraper_api.py)."""
    from datetime import date, timedelta
    from collections import defaultdict
    heute = date.today()
    bis   = heute + timedelta(days=days_ahead)

    praxis = next((p for p in PRAXEN_DATA if p["entity_id"] == practice_id), None)
    therapeuten = praxis["therapeuten"] if praxis else []
    if prof_id:
        therapeuten = [t for t in therapeuten if t["prof_id"] == prof_id]

    alle_slots = []
    for th in therapeuten[:8]:
        apt_type_id = await get_apt_type_id(session, practice_id, th["prof_id"])
        status, data = await onedoc_get(
            session,
            f"/api/v1/locations/{practice_id}/availabilities",
            {
                "professionId": "53",
                "startDate": heute.isoformat(),
                "endDate": bis.isoformat(),
                "professionalId": th["prof_id"],
                "appointmentTypeId": apt_type_id,
                "acceptedClientType": "established",
                "selectUniqueAppointmentTypePerDay": "true",
                "field": "availabilities:timeSlots,nextTimeSlotOn",
            },
        )
        if status == 200 and data:
            time_slots = data.get("data", {}).get("timeSlots", {})
            for datum, slot_list in time_slots.items():
                for ts in slot_list:
                    alle_slots.append({
                        "datum": datum,
                        "zeit": ts.get("dateTime", "")[11:16],
                        "datetime_iso": ts.get("dateTime", ""),
                        "therapeut": th["name"],
                        "prof_id": th["prof_id"],
                        "calendar_id": th.get("calendar_id", ""),
                        "appointment_type_id": apt_type_id,
                        "entity_id": practice_id,
                    })

    by_date: dict = defaultdict(list)
    for s in alle_slots:
        by_date[s["datum"]].append(s)

    freie_tage = []
    for datum in sorted(by_date.keys()):
        slots_day = sorted(by_date[datum], key=lambda x: x["zeit"])
        freie_tage.append({
            "datum": datum,
            "anzahl": len(slots_day),
            "slots": slots_day[:12],
        })

    if not alle_slots:
        print(f"OneDoc: keine Slots für {standort} — möglicherweise Timeout oder keine Verfügbarkeit")

    return {
        "standort":    standort,
        "entity_id":   practice_id,
        "zeitraum":    f"{heute.isoformat()} – {bis.isoformat()}",
        "freie_tage":  freie_tage,
        "total_slots": len(alle_slots),
        "hinweis": "keine Slots" if not alle_slots else None,
    }


SLOTS_PROXY    = os.environ.get("SLOTS_PROXY_URL", "http://91.99.179.44:8098")
SLOTS_API_KEY  = os.environ.get("SLOTS_API_KEY", "")

@app.get("/demo-patient")
async def demo_patient():
    from fastapi.responses import FileResponse
    return FileResponse("demo_patient.html")

@app.get("/demo-team")
async def demo_team():
    from fastapi.responses import FileResponse
    return FileResponse("demo_team.html")

@app.get("/api/hyrox-kurse")
async def get_hyrox_kurse():
    """Live HYROX Stundenplan via SportNow."""
    try:
        headers = {}
        if SLOTS_API_KEY:
            headers["X-API-Key"] = SLOTS_API_KEY
        async with httpx.AsyncClient(timeout=15) as session:
            r = await session.get(f"{SLOTS_PROXY}/hyrox-kurse", headers=headers)
            return JSONResponse(r.json())
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/termine")
async def get_live_termine(standort: str = "", tage: int = 14):
    """Slots via Hetzner-Proxy holen (umgeht Render-Firewall)."""
    try:
        params = f"tage={tage}"
        if standort:
            params += f"&standort={standort}"
        headers = {}
        if SLOTS_API_KEY:
            headers["X-API-Key"] = SLOTS_API_KEY
        async with httpx.AsyncClient(timeout=30) as session:
            r = await session.get(f"{SLOTS_PROXY}/slots?{params}", headers=headers)
            return JSONResponse(r.json())
    except Exception as e:
        return JSONResponse({"error": str(e), "freie_tage": [], "total_slots": 0}, status_code=500)

# Static Files MUSS zuletzt stehen — nach allen API-Routen
app.mount("/", StaticFiles(directory=".", html=True), name="frontend")
